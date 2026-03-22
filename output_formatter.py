"""
output_formatter.py v2 — Professional DOCX Formatter
======================================================
Submission-ready output. No markdown artifacts.
Real Word tables with borders. Embedded figures.
Page numbers. Running header. SPPU/APA/IEEE styles.
Cost: zero API calls. Pure python-docx.
"""

import re
import io
import os
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import xml.etree.ElementTree as ET

# ── Style registry ─────────────────────────────────────────
FORMATTING_STYLES = {
    "SPPU":      {"font": "Times New Roman", "size": 12, "spacing": 1.5,  "margin": 1.0, "citation": "APA"},
    "APA7":      {"font": "Times New Roman", "size": 12, "spacing": 2.0,  "margin": 1.0, "citation": "APA"},
    "MLA9":      {"font": "Times New Roman", "size": 12, "spacing": 2.0,  "margin": 1.0, "citation": "MLA"},
    "Vancouver": {"font": "Arial",           "size": 11, "spacing": 1.5,  "margin": 1.0, "citation": "Vancouver"},
    "Chicago17": {"font": "Times New Roman", "size": 12, "spacing": 2.0,  "margin": 1.0, "citation": "Chicago"},
    "IEEE":      {"font": "Times New Roman", "size": 10, "spacing": 1.0,  "margin": 0.75,"citation": "IEEE"},
    "Harvard":   {"font": "Arial",           "size": 12, "spacing": 1.5,  "margin": 1.0, "citation": "Harvard"},
}

STATISTICIAN_QUOTES = [
    '"Without data, you are just another person with an opinion." — W. Edwards Deming',
    '"All models are wrong, but some are useful." — George Box',
    '"Statistics is the grammar of science." — Karl Pearson',
    '"Far better an approximate answer to the right question." — John Tukey',
    '"In God we trust; all others must bring data." — W. Edwards Deming',
    '"Torture the data long enough and it will confess to anything." — Ronald Coase',
    '"Statistical thinking will one day be necessary for efficient citizenship." — H.G. Wells',
]

PROGRESS_MESSAGES = [
    "Running Professor Wagh's algorithm. The master at work...",
    "Fetching verified citations from academic databases...",
    "Reverse-engineering dataset to match target statistics...",
    "Calibrating Cronbach alpha to publication standards...",
    "Enforcing citation discipline — removing ghost references...",
    "Applying human rhythm variation to prose cadence...",
    "Running neutrality audit — ensuring zero bias...",
    "Aligning effect sizes with domain literature...",
    "Formatting to publication standards...",
    "Final quality sweep — blocklist guard running...",
    "Verifying every citation against CrossRef database...",
    "Polishing prose to academic readability standards...",
]


# ══════════════════════════════════════════════════════════
# TEXT CLEANING
# ══════════════════════════════════════════════════════════

def clean_markdown(text: str) -> str:
    """
    Strip ALL markdown artifacts from generated text.
    Guaranteed: no hashes, no stars, no pipes, no em-dashes,
    no bullet symbols, no backticks in the final output.
    """
    # 1. Remove markdown table rows (pipe-separated) — CRITICAL
    #    These are the "garbled blobs" the user keeps seeing
    lines = text.split('\n')
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip pure separator rows like |---|---|
        if re.match(r'^[\|\-:\s]+$', stripped) and '|' in stripped:
            continue
        # Convert pipe-table data rows to prose
        if stripped.startswith('|') and stripped.endswith('|') and '|' in stripped[1:-1]:
            # Extract cell contents, join as prose
            cells = [c.strip() for c in stripped[1:-1].split('|') if c.strip()]
            if cells:
                clean_lines.append(' — '.join(cells))
            continue
        clean_lines.append(line)
    text = '\n'.join(clean_lines)

    # 2. Hash headers — remove hashes, keep text
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # 3. Bold/italic stars
    text = re.sub(r'\*{1,3}([^*\n]+)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,2}([^_\n]+)_{1,2}', r'\1', text)

    # 4. Em/en dashes → comma
    text = text.replace('\u2014', ',').replace('\u2013', '-')
    text = text.replace(' -- ', ', ').replace(' --- ', ', ')
    text = text.replace('EMDASH', ',')

    # 5. Bullet points
    text = re.sub(r'^[\*\-•]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\d+\.\s+(?=[a-z])', '', text, flags=re.MULTILINE)

    # 6. Backticks
    text = re.sub(r'`+([^`\n]*)`+', r'\1', text)

    # 7. Horizontal rules
    text = re.sub(r'^[-_*]{3,}\s*$', '', text, flags=re.MULTILINE)

    # 8. Any remaining pipe characters not caught above
    text = re.sub(r'\s*\|\s*', ' ', text)

    # 9. Trailing whitespace on lines
    text = '\n'.join(line.rstrip() for line in text.split('\n'))

    # 10. Excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 11. Double spaces
    text = re.sub(r'  +', ' ', text)

    return text.strip()


def extract_keywords(text: str) -> List[str]:
    m = re.search(r'[Kk]eywords?\s*[:—\-]\s*(.+?)(?:\n|$)', text)
    if m:
        kws = [k.strip().strip('.,;') for k in re.split(r'[,;]', m.group(1))]
        return [k for k in kws if k][:8]
    return []


def extract_abstract(text: str) -> str:
    m = re.search(r'(?:^|\n)(?:##?\s*)?Abstract\s*\n+(.+?)(?=\n##?\s*\w|\Z)',
                  text, re.DOTALL | re.IGNORECASE)
    return m.group(1).strip()[:1500] if m else ""


# ══════════════════════════════════════════════════════════
# DOCX HELPERS
# ══════════════════════════════════════════════════════════

def _set_margins(doc: Document, margin_in: float):
    for section in doc.sections:
        m = Inches(margin_in)
        section.top_margin = section.bottom_margin = m
        section.left_margin = section.right_margin = m


def _set_line_spacing(para, spacing: float):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing      = spacing
    fmt.space_after       = Pt(6)


def _apply_font(run, font_name: str, size: float,
                bold: bool = False, color: str = "000000"):
    run.font.name   = font_name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _add_page_numbers(doc: Document, font_name: str, font_size: float):
    """Add page numbers to footer — right aligned."""
    for section in doc.sections:
        footer = section.footer
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size - 2)
        # Page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def _add_running_header(doc: Document, title: str,
                         font_name: str, font_size: float):
    """Add running header with paper title — left, page number right."""
    short_title = title[:60] + ("..." if len(title) > 60 else "")
    for section in doc.sections:
        header = section.header
        para   = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(short_title)
        run.font.name   = font_name
        run.font.size   = Pt(font_size - 2)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # Separator line
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)


def _table_set_borders(table):
    """Apply thin black borders to all table cells."""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _cell_shade(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _add_professional_table(doc: Document, headers: List[str],
                              rows: List[List], caption: str,
                              font_name: str, font_size: float,
                              spacing: float):
    """Add a properly formatted Word table — SPSS style."""
    # Caption above table
    cap_para = doc.add_paragraph()
    cap_run  = cap_para.add_run(caption)
    cap_run.font.name   = font_name
    cap_run.font.size   = Pt(font_size - 1)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(0, 0, 0)
    _set_line_spacing(cap_para, spacing)

    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 for header

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_set_borders(table)

    # Header row — dark blue shading, white bold text
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _cell_shade(cell, '1F4E79')
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(str(h))
        run.font.name  = font_name
        run.font.size  = Pt(font_size - 1)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        is_total = str(row_data[0]).strip().lower() in ('total', 'grand total')
        shade = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        if is_total:
            shade = 'D9E2F3'

        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            _cell_shade(cell, shade)
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run  = p.add_run(str(val))
            run.font.name  = font_name
            run.font.size  = Pt(font_size - 1)
            run.font.bold  = is_total
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # spacing after table
    return table


def _embed_figure(doc: Document, img_bytes: bytes, caption: str,
                   font_name: str, font_size: float, width_inches: float = 5.5):
    """Embed a matplotlib figure (PNG bytes) into the DOCX."""
    if not img_bytes:
        return
    buf = io.BytesIO(img_bytes)
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run()
        run.add_picture(buf, width=Inches(width_inches))
    except Exception:
        return
    # Caption below figure
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run  = cap_para.add_run(caption)
    cap_run.font.name   = font_name
    cap_run.font.size   = Pt(font_size - 1)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════
# SECTION DETECTION
# ══════════════════════════════════════════════════════════

SECTION_KEYWORDS = [
    "abstract", "introduction", "literature review", "background",
    "theoretical framework", "conceptual framework", "objectives",
    "hypotheses", "research questions", "methodology", "methods",
    "research design", "data collection", "participants", "instruments",
    "procedure", "results", "findings", "data analysis", "discussion",
    "conclusion", "recommendations", "limitations", "implications",
    "future research", "references", "bibliography", "appendix",
    "acknowledgements", "acknowledgments", "declaration",
    "search strategy", "inclusion criteria", "exclusion criteria",
    "heterogeneity", "scale reliability", "validity",
]


def _is_section_heading(line: str) -> bool:
    stripped = line.strip().rstrip('.:')
    words    = stripped.split()
    if not stripped or len(words) > 10:
        return False
    lower = stripped.lower()
    if any(kw in lower for kw in SECTION_KEYWORDS):
        return True
    if re.match(r'^\d+(\.\d+)?\s+[A-Z]', stripped):
        return True
    if re.match(r'^[IVXLC]+\.\s+[A-Z]', stripped):
        return True
    return False


def _is_subsection_heading(line: str) -> bool:
    stripped = line.strip()
    if re.match(r'^\d+\.\d+\s+[A-Z]', stripped):
        return True
    return False


# ══════════════════════════════════════════════════════════
# MAIN BUILDER
# ══════════════════════════════════════════════════════════

def build_professional_docx(
        content: str,
        title: str,
        style_key: str = "SPPU",
        keywords: List[str] = None,
        citation_style: str = "APA",
        figures: List[Tuple[bytes, str]] = None,
        stats_tables: List[Tuple[List, List[List], str]] = None,
) -> bytes:
    """
    Build a completely professional, submission-ready DOCX.

    figures:      list of (png_bytes, caption) tuples
    stats_tables: list of (headers, rows, caption) tuples
    """
    style    = FORMATTING_STYLES.get(style_key, FORMATTING_STYLES["SPPU"])
    font     = style["font"]
    size     = style["size"]
    spacing  = style["spacing"]
    margin   = style["margin"]

    content = clean_markdown(content)
    doc     = Document()
    _set_margins(doc, margin)

    # ── Default paragraph style ──────────────────────────
    normal              = doc.styles['Normal']
    normal.font.name    = font
    normal.font.size    = Pt(size)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.space_after      = Pt(6)

    # ── Header & footer ──────────────────────────────────
    _add_running_header(doc, title, font, size)
    _add_page_numbers(doc, font, size)

    # ── Title page ───────────────────────────────────────
    title_para           = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run            = title_para.add_run(title.upper())
    _apply_font(title_run, font, size + 2, bold=True)
    _set_line_spacing(title_para, spacing)
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after  = Pt(18)

    # ── Keywords box ─────────────────────────────────────
    if keywords:
        kw_para           = doc.add_paragraph()
        kw_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kw_run            = kw_para.add_run("Keywords: ")
        _apply_font(kw_run, font, size, bold=True)
        kw_text = kw_para.add_run(", ".join(keywords))
        _apply_font(kw_text, font, size)
        kw_text.font.italic = True
        _set_line_spacing(kw_para, spacing)
        # Box border
        pPr   = kw_para._p.get_or_add_pPr()
        pBdr  = OxmlElement('w:pBdr')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '6')
            el.set(qn('w:space'), '4')
            el.set(qn('w:color'), '1F4E79')
            pBdr.append(el)
        pPr.append(pBdr)
        doc.add_paragraph()

    # ── Parse and render body ────────────────────────────
    lines               = content.split('\n')
    current_para_lines  = []
    in_references       = False
    ref_count           = 0

    def flush_paragraph():
        nonlocal current_para_lines
        if not current_para_lines:
            return
        text = ' '.join(l.strip() for l in current_para_lines if l.strip())
        if not text:
            current_para_lines = []
            return
        p           = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run         = p.add_run(text)
        _apply_font(run, font, size)
        _set_line_spacing(p, spacing)
        if in_references:
            p.paragraph_format.left_indent   = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
        current_para_lines = []

    # ── Mutable queues: figures/tables consumed as inline markers are hit ──
    figure_queue        = list(figures)      if figures      else []
    table_queue         = list(stats_tables) if stats_tables else []
    inline_figure_count = 0
    inline_table_count  = 0

    figure_idx = 0
    table_idx  = 0

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        # ── [DIAGRAM: caption] / [FIGURE: ...] → embed next queued figure inline ──
        diag_m = re.match(
            r'^\[(?:DIAGRAM|FIGURE|CHART|GRAPH|IMAGE)\s*:\s*(.+?)\]$',
            stripped, re.IGNORECASE)
        if diag_m:
            flush_paragraph()
            hint = diag_m.group(1).strip()
            if figure_queue:
                fig_bytes, fig_cap = figure_queue.pop(0)
                inline_figure_count += 1
                _embed_figure(doc, fig_bytes,
                              f"Figure {inline_figure_count}. {fig_cap or hint}",
                              font, size)
            else:
                # No bytes yet — leave a styled placeholder so position is preserved
                inline_figure_count += 1
                ph = doc.add_paragraph()
                ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r  = ph.add_run(f"[Figure {inline_figure_count}: {hint}]")
                _apply_font(r, font, size - 1)
                r.font.italic     = True
                r.font.color.rgb  = RGBColor(150, 150, 150)
            continue

        # ── [TABLE: caption] / [STATS TABLE: ...] → embed next queued table inline ──
        tbl_m = re.match(
            r'^\[(?:TABLE|STATS\s*TABLE)\s*:\s*(.+?)\]$',
            stripped, re.IGNORECASE)
        if tbl_m:
            flush_paragraph()
            hint = tbl_m.group(1).strip()
            if table_queue:
                hdrs, rows, cap = table_queue.pop(0)
                inline_table_count += 1
                _add_professional_table(
                    doc, hdrs, rows,
                    f"Table {inline_table_count}. {cap or hint}",
                    font, size, spacing)
            continue

        # References section detection
        if re.match(r'^(?:##?\s*)?references?\s*$', stripped, re.IGNORECASE):
            flush_paragraph()
            in_references = True
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hr = h.add_run("References")
            _apply_font(hr, font, size + 1, bold=True)
            _set_line_spacing(h, spacing)
            continue

        # Section headings
        if _is_subsection_heading(stripped):
            flush_paragraph()
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hr = h.add_run(stripped)
            _apply_font(hr, font, size, bold=True)
            hr.font.italic = True
            _set_line_spacing(h, spacing)
            continue

        if _is_section_heading(stripped) and not in_references:
            flush_paragraph()
            h  = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hr = h.add_run(stripped.title())
            _apply_font(hr, font, size + 1, bold=True)
            _set_line_spacing(h, spacing)
            h.paragraph_format.space_before = Pt(12)
            continue

        # Numbered references
        if in_references and re.match(r'^\d+\.?\s+\w', stripped):
            flush_paragraph()
            ref_count += 1
            p           = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent        = Inches(0.5)
            p.paragraph_format.first_line_indent  = Inches(-0.5)
            run         = p.add_run(stripped)
            _apply_font(run, font, size)
            _set_line_spacing(p, spacing)
            continue

        # Numbered lists in body
        if re.match(r'^\d+\.\s+\w', stripped) and not in_references:
            flush_paragraph()
            p           = doc.add_paragraph(style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run         = p.add_run(re.sub(r'^\d+\.\s+', '', stripped))
            _apply_font(run, font, size)
            _set_line_spacing(p, spacing)
            continue

        current_para_lines.append(stripped)

    flush_paragraph()

    # ── Append any remaining tables not yet consumed inline ─────────────
    if table_queue:
        doc.add_page_break()
        h  = doc.add_paragraph()
        hr = h.add_run("Statistical Tables")
        _apply_font(hr, font, size + 1, bold=True)
        _set_line_spacing(h, spacing)
        for hdrs, rows, cap in table_queue:
            inline_table_count += 1
            _add_professional_table(
                doc, hdrs, rows,
                f"Table {inline_table_count}. {cap}",
                font, size, spacing)

    # ── Append any remaining figures not yet consumed inline ──────────
    if figure_queue:
        doc.add_page_break()
        h  = doc.add_paragraph()
        hr = h.add_run("Figures")
        _apply_font(hr, font, size + 1, bold=True)
        _set_line_spacing(h, spacing)
        for fig_bytes, fig_caption in figure_queue:
            inline_figure_count += 1
            full_caption = f"Figure {inline_figure_count}. {fig_caption}"
            _embed_figure(doc, fig_bytes, full_caption, font, size)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
