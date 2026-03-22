"""
PaperForge AI — Universal Academic Paper Simulator
===================================================
15-step pipeline. Human in loop at key decision points.
Credits enforced server-side via Supabase.
"""

import streamlit as st
import anthropic, os, json, re, io, uuid, time
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile

# ── Local modules ──────────────────────────────────────────
from model_router    import call_model, call_prep, call_writer, call_audit, WRITER_MODELS
from citation_engine import (fetch_citation_bank, enforce_citation_discipline,
                              verify_bank_phase_c, format_citation, bank_to_prompt_text)
from data_engine     import generate_narrative_targets, reverse_engineer_dataset, verify_statistics
from audit_pipeline  import audit_and_clean

from credit_engine   import get_engine, CREDIT_COSTS
from prompt_bank     import get_bank
try:
    from msata_component    import render_msata_step
    from diagram_engine     import (generate_figures_for_paper, generate_chart,
                                    decide_chart_type)
    from likert_engine      import (render_construct_editor, generate_constructs_prompt,
                                    generate_spss_frequency_table, generate_crosstab,
                                    suggest_crosstabs, MIN_ITEMS_INDIA)
    from otp_engine         import render_otp_verification
    from conference_matcher import (extract_conference_themes,
                                    match_domain_to_themes, suggest_paper_title)
    EXTENSIONS_AVAILABLE = True
except ImportError as _ie:
    EXTENSIONS_AVAILABLE = False
    print(f"Extensions not loaded: {_ie}")
    MSATA_AVAILABLE = True
except ImportError:
    MSATA_AVAILABLE = False
try:
    from alacarte_parser import parse_alacarte, format_validation_report, FORMATTING_STYLES, STATISTICIAN_QUOTES, PROGRESS_MESSAGES
    from output_formatter import build_professional_docx, clean_markdown, extract_keywords
    FORMATTER_AVAILABLE = True
except ImportError:
    FORMATTER_AVAILABLE = False
    STATISTICIAN_QUOTES = ['"Without data, you are just another person with an opinion." — W. Edwards Deming']
    PROGRESS_MESSAGES = ["Running Professor Wagh's algorithm. The master at work..."]

# ══════════════════════════════════════════════════════════
# PAGE CONFIG & CSS
# ══════════════════════════════════════════════════════════
st.set_page_config(page_title="PaperForge AI", page_icon="📄",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;800&family=Source+Serif+4:wght@300;400;600&display=swap');
html, body, [class*="css"] { font-family: 'Source Serif 4', Georgia, serif !important; background: #F5F2ED !important; color: #1a1a1a !important; font-size: 16px !important; }
h1 { font-family: 'Playfair Display', serif !important; font-size: 2.4rem !important; color: #1a1a1a !important; font-weight: 800 !important; }
h2 { font-family: 'Playfair Display', serif !important; font-size: 1.6rem !important; color: #1a1a1a !important; }
h3 { font-family: 'Source Serif 4', serif !important; font-size: 1.2rem !important; color: #2a2a2a !important; font-weight: 600 !important; }
p, div, span, label { color: #1a1a1a !important; font-size: 15px !important; }
.stButton > button { background: #8B4513 !important; color: #FFFFFF !important; border: none !important; border-radius: 3px !important; font-family: 'Source Serif 4', serif !important; font-weight: 600 !important; font-size: 15px !important; padding: 0.6rem 1.8rem !important; }
.stButton > button:hover { background: #6B3410 !important; }
.step-box { background: #FFFFFF; border-left: 4px solid #8B4513; padding: 1.2rem 1.5rem; border-radius: 3px; margin: 0.8rem 0; color: #1a1a1a !important; box-shadow: 0 1px 4px rgba(0,0,0,0.08); }
.warn-box { background: #FFF8E7; border-left: 4px solid #D4820A; padding: 1rem 1.2rem; border-radius: 3px; margin: 0.5rem 0; color: #5a3800 !important; font-size: 14px !important; }
.ok-box { background: #F0FFF4; border-left: 4px solid #2D7A3A; padding: 1rem 1.2rem; border-radius: 3px; margin: 0.5rem 0; color: #1a3a1a !important; }
.narrative-card { background: #FFFFFF; border: 2px solid #D4C5B0; border-radius: 6px; padding: 1.5rem; margin: 0.8rem 0; color: #1a1a1a !important; box-shadow: 0 2px 8px rgba(0,0,0,0.06); }
.narrative-selected { border-color: #8B4513 !important; background: #FFF8F0 !important; }
.credit-drain { color: #CC0000 !important; font-size: 13px !important; font-weight: 700 !important; animation: flash 0.4s ease-in-out 3; }
@keyframes flash { 0%,100%{opacity:1} 50%{opacity:0.2} }
.quote-box { background: #FFF8F0; border: 1px solid #D4C5B0; border-radius: 4px; padding: 1rem 1.5rem; margin: 1rem 0; font-style: italic; color: #4a3520 !important; font-size: 14px !important; }
div[data-testid="stSidebarContent"] { background: #FFFFFF !important; border-right: 1px solid #D4C5B0 !important; }
div[data-testid="stSidebarContent"] * { color: #1a1a1a !important; }
.stTabs [data-baseweb="tab"] { background: #F5F2ED !important; color: #4a3520 !important; font-size: 14px !important; }
.stTabs [aria-selected="true"] { background: #8B4513 !important; color: #FFFFFF !important; font-weight: 600 !important; }
.stTextArea textarea, .stTextInput input { background: #FFFFFF !important; color: #1a1a1a !important; border: 1px solid #C5B9A8 !important; font-family: 'Source Serif 4', serif !important; font-size: 15px !important; }
.stProgress > div > div { background: #8B4513 !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════
DEFAULTS = {
    "step": 1,
    "paper_id": str(uuid.uuid4()),
    "user_id": "tester_001",  # replace with auth in prod
    "topic": "",
    "paper_type": "Research Paper",
    "tier": "Basic",
    "word_limit": 4000,
    "citation_style": "APA",
    "language": "English",
    "uploaded_material": "",
    "conference_rules": "",
    "conference_constraints": {},
    "domain_analysis": {},
    "template_key": "Generic_IMRaD",
    "citation_bank": [],
    "objectives": [],
    "hypotheses": [],
    "structure": [],
    "stats_plan": {},
    "selected_narrative": None,
    "narratives": {},
    "target_stats": {},
    "synthetic_df": None,
    "stats_verification": {},
    "full_paper": "",
    "audit_issues": [],
    "user_overrides": "",
    "regen_count_local": 0,
    "msata_signed": False,
    "coupon_validated": False,
    "coupon_code": "",
    "feedback_given": False,
    "constructs": [],
    "selected_crosstabs": [],
    "table_mode": "combined",
    "show_diagrams": True,
    "language_complexity": 11,
    "sample_size_override": None,
    "conference_themes": {},
    "conference_match": [],
    "otp_verified": False,
    "fetch_more_count": 0,
    "include_diagrams": False,
    "stats_level": "Basic",
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

credits = get_engine()
bank    = get_bank()

# ══════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════

def get_api_keys() -> dict:
    return {
        "anthropic": os.environ.get("ANTHROPIC_API_KEY", st.session_state.get("api_key_input","")),
        "fireworks": os.environ.get("FIREWORKS_API_KEY",""),
        "groq":      os.environ.get("GROQ_API_KEY",""),
        "deepseek":  os.environ.get("DEEPSEEK_API_KEY",""),
    }

def claude_call(system: str, user: str, max_tokens: int = 3000,
                heavy: bool = False, cheap: bool = False) -> str:
    """
    cheap=True  → call_prep() → DeepSeek→Groq→Haiku  (~₹0.05/call)
    cheap=False → call_writer() → tier-appropriate model (~₹4-45/call)
    heavy=True  → Opus for Premium/Ultra heavy sections
    """
    api_keys = get_api_keys()
    msgs = [{"role": "user", "content": user}]
    if cheap:
        return call_prep(msgs, system, max_tokens, api_keys)
    return call_writer(st.session_state.tier, msgs, system,
                       max_tokens, heavy, api_keys)

def json_parse(text: str) -> any:
    clean = re.sub(r'```json|```','',text).strip()
    return json.loads(clean)

def word_count(text: str) -> int:
    return len(re.findall(r'\b\w+\b', text))

def show_credit_drain(cost: float, label: str = ""):
    if cost > 0:
        st.markdown(f'<div class="credit-drain">−{cost} credit {"" if not label else "· "+label}</div>',
                    unsafe_allow_html=True)

def build_docx(content: str, title: str) -> bytes:
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

MASTER_SYSTEM = """You are PaperForge Oracle — a world-class academic writing engine.
Universal: any domain, any paper type, any field.

ABSOLUTE RULES — VIOLATION IS NOT PERMITTED:
1. UPLOADED MATERIAL IS LAW — every stat/figure/citation from uploads MUST appear.
2. NEVER invent citations. Only cite from the provided bank. Zero exceptions.
3. Honor every user override without exception.
4. OUTPUT FORMAT: Plain prose only. Use ## for section headings, ### for subsections.
5. NEVER use: em-dashes (—), en-dashes (–), double-hyphens (--), ellipsis (...).
6. NEVER use pipe characters (|) in prose. Tables are forbidden in your output.
   The formatter handles all tables. You write prose only.
7. NEVER use markdown bullets (*, -, •) inside paragraphs.
8. NEVER use bold (**text**) or italic (*text*) in body paragraphs.
9. NEVER use hash symbols (#) except at the very start of a section heading line.
10. Do not use: "simulation", "synthetic", "fabricated", "generated data", "illustrative".
11. Write in complete paragraphs. Every section must flow as continuous prose.
12. Statistics are reported inline: t(119) = 4.50, p = 0.000, Cohen d = 0.41.
13. Consumer is king. Quality first. Every word earns its place."""


def _build_stats_text(v: dict, narrative: dict) -> str:
    lines = []
    for hv in v.get("hypotheses",[]):
        idx = hv["hypothesis_num"] - 1
        ht = narrative.get("hypotheses",[{}]*3)
        if idx < len(ht):
            h = ht[idx]
            direction = "supported" if h.get("supported") else "not supported"
            lines.append(
                f"H{hv['hypothesis_num']}: t({narrative.get('n',60)-1}) = {hv.get('t_statistic','-')}, "
                f"p = {hv.get('p_value','-')}, Cohen's d = {hv.get('cohens_d','-')} "
                f"[{direction}, {h.get('effect_label','')} effect]"
            )
    ca = v.get("cronbach_alpha",{})
    if ca.get("computed"):
        lines.append(f"Cronbach's alpha = {ca['computed']} (scale reliability)")
    return "\n".join(lines)


def _replace_section(full_paper: str, section_name: str, new_content: str) -> str:
    """Best-effort section replacement."""
    pattern = rf'(##\s+{re.escape(section_name)}.*?)(?=\n##\s|\Z)'
    replacement = f"## {section_name}\n\n{new_content}"
    result = re.sub(pattern, replacement, full_paper, flags=re.DOTALL | re.IGNORECASE)
    return result if result != full_paper else full_paper + f"\n\n## {section_name}\n\n{new_content}"


# ══════════════════════════════════════════════════════════
# SECTIONAL PAPER GENERATOR — fixes truncation at ~1,400w
# ══════════════════════════════════════════════════════════
# Root cause: Haiku (Basic tier) reliably outputs only ~1,400
# words per call when given an 8,000-token budget. The model
# self-truncates before filling its context window.
# Fix: generate each section in its own API call (2,200 tok
# budget each), then concatenate. Works on any model tier.
# ══════════════════════════════════════════════════════════

def _section_tokens(word_allocation: int) -> int:
    """Convert target word count → safe max_tokens for one section call."""
    # 1 word ≈ 1.3 tokens. Add 20% headroom, floor at 800, cap at 2800.
    return min(2800, max(800, int(word_allocation * 1.3 * 1.20)))


def generate_paper_sectional(
    topic: str,
    paper_type: str,
    tier: str,
    domain: str,
    language: str,
    grade: int,
    tone: str,
    structure: list,
    stat_results: str,
    objectives: list,
    hypotheses: list,
    cites_text: str,
    uploaded_material: str,
    conference_constraints: dict,
    user_overrides: str,
    citation_style: str,
    stats_level: str,
    include_diagrams: bool,
    citation_bank: list,
    progress_bar,
    progress_msg,
) -> str:
    """
    Write each section as a separate API call, then stitch.

    Benefits vs single-call approach:
    - No truncation: each call has a focused, achievable budget.
    - Better coherence: each section prompt includes what came before.
    - Works on Haiku (Basic) and Sonnet (Medium+) equally.
    - Progress bar reflects real section-by-section completion.
    """
    diagram_instr = (
        "If appropriate, mark diagram positions as [DIAGRAM: brief description]."
        if include_diagrams else "No diagrams required."
    )

    # ── shared context block (sent with every section call) ──
    shared_ctx = f"""PAPER TOPIC: "{topic}"
PAPER TYPE: {paper_type}
DOMAIN: {domain}
LANGUAGE: {language} | GRADE: {grade} (Flesch-Kincaid)
TONE: {tone}
CITATION STYLE: {citation_style}
STATS LEVEL: {stats_level}

VERIFIED STATISTICS (report exactly as shown):
{stat_results or 'None'}

OBJECTIVES:
{json.dumps([o.get("text","") for o in objectives], indent=2)}

HYPOTHESES:
{json.dumps([h.get("alternate","") for h in hypotheses], indent=2)}

REAL CITATIONS (use ONLY these):
{cites_text}

UPLOADED MATERIAL (every stat/figure must appear):
{(uploaded_material or "None")[:2000]}

USER OVERRIDES (absolute law):
{user_overrides or "None"}

CONFERENCE CONSTRAINTS:
{json.dumps(conference_constraints) if conference_constraints else "None"}"""

    # ── section ordering metadata ──
    total_sections = len(structure)
    sections_written: list[str] = []      # accumulates formatted sections
    paper_so_far_words = 0

    for idx, sec in enumerate(structure):
        sec_name  = sec.get("section", f"Section {idx+1}")
        sec_words = int(sec.get("word_allocation") or 400)
        sec_notes = sec.get("notes", "")
        max_tok   = _section_tokens(sec_words)

        pct = int(10 + 80 * idx / total_sections)
        progress_bar.progress(pct)
        progress_msg.markdown(f"⚙️ Writing **{sec_name}** ({sec_words} words)…")

        # ── running context: last 400 words of paper so far ──
        running_tail = ""
        if sections_written:
            all_so_far = "\n\n".join(sections_written)
            tail_words = all_so_far.split()[-400:]
            running_tail = (
                "\n\nPAPER SO FAR (last 400 words — maintain continuity):\n"
                + " ".join(tail_words)
            )

        is_abstract   = "abstract"   in sec_name.lower()
        is_references = "reference"  in sec_name.lower() or "bibliography" in sec_name.lower()

        if is_references:
            # References: just format the citation bank, no AI call needed
            ref_lines = []
            for i, c in enumerate(citation_bank[:20], 1):
                authors = c.get("authors", ["Unknown"])
                year    = c.get("year", "n.d.")
                title   = c.get("title", "Untitled")
                journal = c.get("journal", "")
                doi     = c.get("doi", "")
                if citation_style.upper().startswith("APA"):
                    auth_str = ", ".join(authors[:3]) + (" et al." if len(authors) > 3 else "")
                    line = f"{auth_str} ({year}). {title}."
                    if journal: line += f" *{journal}*."
                    if doi:     line += f" https://doi.org/{doi}"
                else:
                    line = f"[{i}] " + "; ".join(authors[:3]) + f" ({year}). {title}."
                    if journal: line += f" {journal}."
                ref_lines.append(line)
            section_text = "## References\n\n" + "\n\n".join(ref_lines)
            sections_written.append(section_text)
            continue

        section_prompt = f"""{shared_ctx}{running_tail}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION TO WRITE NOW: {sec_name}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Target length: EXACTLY {sec_words} words (±5%).
Section notes: {sec_notes or "Standard academic prose."}
Section position: {idx+1} of {total_sections}.
{diagram_instr}

CRITICAL OUTPUT RULES:
- Start with: ## {sec_name}
- Write ONLY this section. Do not write any other section.
- MINIMUM {int(sec_words * 0.92)} words. Count carefully.
- No markdown bullets, bold, italic, or pipe characters.
- No em-dashes, en-dashes, ellipsis.
- Inline statistics: t(119)=4.50, p=0.000, Cohen d=0.41
- {"Structured format: Background / Methods / Results / Conclusion (one paragraph each)." if is_abstract else "Continuous prose paragraphs only. No sub-lists."}
- Do NOT repeat content from previous sections.
- End the section cleanly. Stop immediately after.

Write ## {sec_name} now:"""

        try:
            section_text = claude_call(MASTER_SYSTEM, section_prompt,
                                       max_tokens=max_tok, heavy=False, cheap=False)
            # Ensure heading is present
            if not section_text.strip().startswith("##"):
                section_text = f"## {sec_name}\n\n{section_text.strip()}"
        except Exception as e:
            section_text = f"## {sec_name}\n\n[Generation error for this section: {e}]"

        sections_written.append(section_text.strip())
        paper_so_far_words += word_count(section_text)

    full_paper = "\n\n".join(sections_written)
    return full_paper


# ══════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## 📄 PaperForge AI")
    st.markdown("---")

    if not os.environ.get("ANTHROPIC_API_KEY"):
        k = st.text_input("Anthropic API Key", type="password", key="api_key_input")

    st.markdown("### Tier")
    tier = st.selectbox("", ["Basic (₹999)", "Medium (₹2,000)", "Advanced (₹5,000)"], label_visibility="collapsed")
    st.session_state.tier = tier.split(" ")[0]
    st.session_state.word_limit = {"Basic":4000,"Medium":5500,"Advanced":7000}[st.session_state.tier]

    st.markdown("### Citation Style")
    st.session_state.citation_style = st.selectbox("", ["APA","Vancouver","MLA"], label_visibility="collapsed")

    st.markdown("### Language")
    lang = st.selectbox("", ["English","Marathi"], label_visibility="collapsed")
    st.session_state.language = lang

    balance = float(st.session_state.get("live_balance", 9999.0))
    warn = None  # Using live_balance system — see sidebar
    # RPG-style credit display
    basic_papers  = int(balance // 104)
    medium_papers = int(balance // 160)
    adv_papers    = int(balance // 240)
    
    if balance >= 500:
        credit_color = "#2D7A3A"   # green — healthy
    elif balance >= 20:
        credit_color = "#D4820A"   # amber — low
    else:
        credit_color = "#CC0000"   # red — critical
    
    st.markdown(f'<div style="background:#f9f9f9;border:1px solid #C5B9A8;border-radius:4px;padding:0.7rem;margin:0.3rem 0">'
                f'<div style="font-size:0.75rem;color:#666;font-weight:600;text-transform:uppercase;letter-spacing:1px">Credits</div>'
                f'<div style="font-size:1.8rem;font-weight:800;color:{credit_color};line-height:1.1">{balance:.0f}</div>'
                f'<div style="font-size:0.72rem;color:#888;margin-top:2px">'
                f'≈ {basic_papers}B / {medium_papers}M / {adv_papers}A papers</div>'
                f'</div>', unsafe_allow_html=True)
    
    if balance < 50:
        st.markdown('<div style="background:#FFF0F0;border-left:3px solid #CC0000;padding:0.4rem 0.6rem;font-size:0.78rem;color:#CC0000;border-radius:2px">🔴 Critical — top up now</div>', unsafe_allow_html=True)
    elif balance < 150:
        st.markdown(f'<div style="background:#FFF8E7;border-left:3px solid #D4820A;padding:0.4rem 0.6rem;font-size:0.78rem;color:#8B5E00;border-radius:2px">⚠️ Low — less than 2 papers left</div>', unsafe_allow_html=True)
    if warn:
        st.markdown(f'<div class="warn-box">{warn}</div>', unsafe_allow_html=True)

    st.markdown(f"**Step:** {st.session_state.step} / 15")
    st.progress(st.session_state.step / 15)

    st.markdown("### Formatting Style")
formatting_options = ["SPPU (Default)", "APA7", "MLA9", "Vancouver", "Chicago17", "IEEE", "Harvard"]
fmt_choice = st.selectbox("", formatting_options, label_visibility="collapsed")
st.session_state["formatting_style"] = fmt_choice.split(" ")[0]

if st.button("🔄 Full Reset"):
        for k2, v2 in DEFAULTS.items():
            st.session_state[k2] = v2
        st.session_state.paper_id = str(uuid.uuid4())
        st.rerun()

# ══════════════════════════════════════════════════════════
# HEADER + STEP INDICATOR
# ══════════════════════════════════════════════════════════
st.markdown("# PaperForge AI")
st.markdown("##### Universal Academic Paper Simulator — Any Domain, Any Field")
st.markdown("---")

STEP_NAMES = [
    "Topic","Domain","Citations","Objectives","Structure",
    "Stats Plan","Narratives","Data","Generate","Audit",
    "Review","MSATA","Download","","",
]
cols = st.columns(min(st.session_state.step + 2, 13))
for i, name in enumerate(STEP_NAMES[:13], 1):
    mark = "🟡" if i==st.session_state.step else ("✅" if i<st.session_state.step else "⚪")
    if i <= len(cols):
        cols[i-1].markdown(f"**{mark}**<br><small>{name}</small>", unsafe_allow_html=True)

st.markdown("---")

# ══════════════════════════════════════════════════════════
# STEP 1 — TOPIC & UPLOADS
# ══════════════════════════════════════════════════════════
# ── COUPON GATE ──────────────────────────────────────────
# ── Credit pack definitions ──────────────────────────────
# 1 credit = ₹5. 10x margin baked in.
CREDIT_PACKS = {
    "starter":     {"credits": 200,  "price": 999,   "label": "Starter",
                    "papers": "2 Basic papers",
                    "features": ["2 Basic papers", "No diagrams", "No CSV", "No supervisor link"],
                    "unlocks": ["basic"]},
    "value":       {"credits": 700,  "price": 2499,  "label": "Value",
                    "papers": "8 Basic / 4 Medium",
                    "features": ["Diagrams included", "CSV download", "8 Basic OR 4 Medium", "Supervisor link"],
                    "unlocks": ["basic", "medium", "diagrams", "csv"]},
    "pro":         {"credits": 2000, "price": 5999,  "label": "Pro",
                    "papers": "20 Basic / 12 Medium / 6 Advanced",
                    "features": ["Everything in Value", "Advanced papers", "Priority routing", "Conference templates"],
                    "unlocks": ["basic", "medium", "advanced", "diagrams", "csv", "conference"]},
    "institution": {"credits": 7000, "price": 14999, "label": "Institution",
                    "papers": "70 Basic / 43 Medium / 20 Advanced",
                    "features": ["Everything", "Bulk coupon generation", "Team access", "Dedicated support"],
                    "unlocks": ["basic", "medium", "advanced", "diagrams", "csv", "conference", "bulk"]},
}

# 1 credit = ₹5. API cost × 10 = credits charged.
TIER_CREDIT_COST = {
    "Basic":    104,
    "Medium":   160,
    "Advanced": 240,
    "Premium":  400,
    "Ultra":    600,
}

ADDON_CREDIT_COST = {
    "section_regen":   16,   # ₹80
    "extra_citations":  5,   # ₹25 — pure profit
    "stats_verify":    10,   # ₹50 — pure profit
    "supervisor":      20,   # ₹100
    "questionnaire":   16,   # ₹80
    "diagrams":         5,   # ₹25 — pure profit
}

VALID_COUPONS = {
    "WAGH9999":   {"credits": 99999, "pack": "institution",
                   "message": "Welcome, Professor Wagh. 99,999 credits. You built this."},
    "TESTER200":  {"credits": 200,   "pack": "starter",
                   "message": "Tester access. 200 credits = 2 Basic papers."},
    "LAUNCH999":  {"credits": 200,   "pack": "starter",
                   "message": "Launch offer! 200 credits = 2 Basic papers free."},
    "PUNE2026":   {"credits": 200,   "pack": "starter",
                   "message": "Pune special. 200 credits = 2 Basic papers."},
    "SPPU2026":   {"credits": 700,   "pack": "value",
                   "message": "SPPU special. 700 credits = 6 Basic papers with diagrams."},
    "EARLYBIRD":  {"credits": 700,   "pack": "value",
                   "message": "Early bird! 700 credits. Diagrams, CSV, everything."},
}

if not st.session_state.get("coupon_validated") and not st.session_state.get("skip_coupon"):
    st.markdown("## Welcome to PaperForge AI")
    st.markdown("Enter your access coupon code to begin, or skip to use available credits.")
    
    c1, c2, c3 = st.columns([3, 1, 1])
    coupon_input = c1.text_input("Coupon Code", placeholder="e.g. PUNE2026", label_visibility="collapsed")
    
    with c2:
        if st.button("Apply Coupon", use_container_width=True):
            code = coupon_input.strip().upper()
            if code in VALID_COUPONS:
                coup = VALID_COUPONS[code]
                st.session_state.coupon_validated = True
                st.session_state.coupon_code = code
                # Add credits
                import re as _re
                # credit loaded via bonus_credits session state
                st.session_state["bonus_credits"] = coup["credits"]
                st.session_state["live_balance"] = coup["credits"]
                st.markdown(f'<div class="ok-box">✅ {coup["message"]}</div>', unsafe_allow_html=True)
                st.rerun()
            else:
                st.error("Invalid coupon code. Try again or skip.")
    with c3:
        if st.button("Skip →", use_container_width=True):
            st.session_state.skip_coupon = True
            st.rerun()
    st.stop()

if st.session_state.step == 1:
    st.markdown("## Step 1 — Topic & Uploads")
    col1, col2 = st.columns([2,1])

    with col1:
        st.session_state.topic = st.text_area(
            "Your topic or title",
            value=st.session_state.topic, height=120,
            placeholder="Any domain. Any field. Be as brief or detailed as you like.\nExamples:\n• Effect of AI tutoring on working memory in LD children\n• Climate change and food security in Sub-Saharan Africa\n• Corporate governance failures in Indian PSU banks")

        paper_types = ["Research Paper","Review Article","Case Study","Meta-Analysis",
                       "Systematic Review","Conference Paper","Thesis Chapter",
                       "Technical Report","Policy Brief","Conceptual Paper"]
        st.session_state.paper_type = st.selectbox("Document type", paper_types)

        col_n1, col_n2 = st.columns([1,2])
        with col_n1:
            n_input = st.text_input("Sample Size (N)", 
                                      value=str(st.session_state.sample_size_override or 90),
                                      placeholder="e.g. 336",
                                      help="Type your exact sample size")
            try:
                st.session_state.sample_size_override = int(n_input.strip())
            except:
                st.session_state.sample_size_override = 90
        with col_n2:
            complexity = st.slider("Language Complexity",
                                    min_value=6, max_value=14,
                                    value=st.session_state.language_complexity,
                                    help="6=Very simple | 10=Standard academic | 12=Research journal | 14=Post-doctoral")
            st.session_state.language_complexity = complexity
            complexity_labels = {6:"Very Simple",7:"Simple",8:"Fairly Simple",9:"Standard",
                                  10:"Academic",11:"Research Journal",12:"Advanced Research",
                                  13:"Professional",14:"Post-Doctoral"}
            st.caption(f"Selected: **{complexity_labels.get(complexity, str(complexity))}**")

        st.session_state.user_overrides = st.text_area(
            "À la carte instructions (your word is absolute law)",
            value=st.session_state.user_overrides, height=70,
            placeholder="e.g. 'Include every figure from upload. Use paired t-test only. Add conceptual framework. Use passive voice throughout.'")

    with col2:
        st.markdown("**Upload reference material**")
        st.caption("All figures, stats, citations found → MANDATORILY included")
        uploaded = st.file_uploader("PDFs, DOCX, CSV, TXT",
                                     accept_multiple_files=True,
                                     type=["pdf","txt","csv","docx"])
        conf_file = st.file_uploader("Conference brochure (optional)",
                                      type=["pdf","txt"])

        if uploaded:
            extracted = []
            for f in uploaded:
                try:
                    if f.name.endswith((".txt",".csv")):
                        extracted.append(f"=== {f.name} ===\n{f.read().decode('utf-8',errors='ignore')}")
                    elif f.name.endswith(".pdf"):
                        import fitz
                        pdf = fitz.open(stream=f.read(), filetype="pdf")
                        text = "\n".join(p.get_text() for p in pdf)
                        extracted.append(f"=== PDF: {f.name} ===\n{text[:8000]}")
                    elif f.name.endswith(".docx"):
                        from docx import Document as DR
                        d = DR(io.BytesIO(f.read()))
                        extracted.append(f"=== DOCX: {f.name} ===\n" +
                                         "\n".join(p.text for p in d.paragraphs)[:8000])
                except Exception as e:
                    extracted.append(f"[Could not read {f.name}: {e}]")
            st.session_state.uploaded_material = "\n\n".join(extracted)
            st.markdown(f'<div class="ok-box">✅ {len(uploaded)} file(s) loaded. All content is mandatory.</div>',
                        unsafe_allow_html=True)

        if conf_file:
            try:
                st.session_state.conference_rules = conf_file.read().decode("utf-8",errors="ignore")[:4000]
                st.markdown('<div class="ok-box">✅ Conference rules loaded.</div>', unsafe_allow_html=True)
            except Exception:
                st.warning("Could not parse conference file.")

    if st.button("▶ Analyse Domain →", use_container_width=True):
        if not st.session_state.topic.strip():
            st.warning("Enter a topic first.")
        else:
            with st.spinner("Analysing topic and domain..."):
                mat = st.session_state.uploaded_material[:3000] if st.session_state.uploaded_material else "None"
                conf = st.session_state.conference_rules[:1000] if st.session_state.conference_rules else "None"
                prompt = f"""Topic: "{st.session_state.topic}"
Document type: {st.session_state.paper_type}
Tier: {st.session_state.tier} ({st.session_state.word_limit} words)
User overrides: {st.session_state.user_overrides or 'None'}
Uploaded material summary: {mat[:1500]}
Conference rules: {conf[:500]}

Return ONLY valid JSON:
{{
  "detected_domain": "...",
  "sub_domain": "...",
  "recommended_methodology": "...",
  "key_variables": ["...","..."],
  "suggested_search_terms": ["...","...","..."],
  "tone_recommendation": "formal|semi-formal|technical|accessible",
  "mandatory_from_upload": ["list stats/figures detected in uploaded material"],
  "conference_hard_rules": {{
    "max_words": null,
    "mandatory_sections": [],
    "abstract_structure": "unstructured",
    "first_person_allowed": false,
    "passive_voice_required": false,
    "reference_format_override": null
  }},
  "domain_explanation": "2-3 sentences"
}}
Return ONLY valid JSON. No markdown fences."""
                try:
                    resp = claude_call(MASTER_SYSTEM, prompt, 2000, cheap=False)
                    st.session_state.domain_analysis = json_parse(resp)
                    # Parse à la carte instructions
                    if st.session_state.user_overrides.strip():
                        parsed = parse_alacarte(
                            st.session_state.user_overrides,
                            tier=st.session_state.tier,
                        ) if FORMATTER_AVAILABLE else None
                        if parsed and (parsed.warnings or parsed.info or parsed.addons):
                            st.session_state["parsed_instructions"] = parsed
                    # Match prompt bank
                    da = st.session_state.domain_analysis
                    key, _ = bank.match(st.session_state.topic, da.get("detected_domain",""))
                    st.session_state.template_key = key
                    # Extract conference constraints
                    cc = da.get("conference_hard_rules", {})
                    if st.session_state.conference_rules:
                        st.session_state.conference_constraints = cc
                    st.session_state.step = 2
                    st.rerun()
                except Exception as e:
                    st.error(f"Domain analysis error: {e}")

# ══════════════════════════════════════════════════════════
# STEP 2 — DOMAIN CONFIRMATION
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 2:
    st.markdown("## Step 2 — Domain Analysis")
    da = st.session_state.domain_analysis

    if da:
        c1,c2,c3 = st.columns(3)
        c1.markdown(f'<div class="step-box"><strong>Domain</strong><br>{da.get("detected_domain","")}<br><small>{da.get("sub_domain","")}</small></div>', unsafe_allow_html=True)
        c2.markdown(f'<div class="step-box"><strong>Methodology</strong><br>{da.get("recommended_methodology","")}</div>', unsafe_allow_html=True)
        c3.markdown(f'<div class="step-box"><strong>Template Matched</strong><br>{st.session_state.template_key}</div>', unsafe_allow_html=True)

        st.markdown(f"**Why:** {da.get('domain_explanation','')}")
    
    # Show à la carte validation
    parsed = st.session_state.get("parsed_instructions")
    if parsed and FORMATTER_AVAILABLE:
        report = format_validation_report(parsed)
        if parsed.warnings:
            st.markdown(f'<div class="warn-box">{report}</div>', unsafe_allow_html=True)
        elif report:
            st.markdown(f'<div class="ok-box">{report}</div>', unsafe_allow_html=True)

        mandatory = da.get("mandatory_from_upload",[])
        if mandatory:
            st.markdown('<div class="ok-box">📌 <strong>Mandatory inclusions from your uploads:</strong><br>' +
                        "<br>".join(f"• {m}" for m in mandatory) + "</div>", unsafe_allow_html=True)

        conf_rules = st.session_state.conference_constraints
        if conf_rules:
            st.markdown("### Conference Hard Rules Detected")
            st.json(conf_rules)

    override = st.text_input("Override domain (leave blank to accept)", placeholder="e.g. Behavioral Neuroscience")
    if override:
        st.session_state.domain_analysis["detected_domain"] = override

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=1; st.rerun()
    with c2:
        if st.button("▶ Fetch Citations →", use_container_width=True):
            st.session_state.step=3; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 3 — CITATION PRE-FETCH (PHASE A)
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 3:
    st.markdown("## Step 3 — Citation Pre-Fetch")
    st.markdown('<div class="step-box">Phase A: Semantic Scholar → OpenAlex → CrossRef cascade. Only verified papers can be cited.</div>', unsafe_allow_html=True)

    da = st.session_state.domain_analysis
    terms = da.get("suggested_search_terms", [st.session_state.topic])

    if not st.session_state.citation_bank:
        with st.spinner("Fetching verified citations..."):
            bank_papers = fetch_citation_bank(terms, target=15)
            st.session_state.citation_bank = bank_papers

    cb = st.session_state.citation_bank
    if cb:
        st.markdown(f"**{len(cb)} papers fetched.**")
        for i, p in enumerate(cb, 1):
            c1,c2 = st.columns([5,1])
            c1.markdown(f"{i}. {format_citation(p, st.session_state.citation_style)}")
            if c2.button("✕", key=f"rem_{i}"):
                st.session_state.citation_bank.pop(i-1); st.rerun()
    else:
        st.markdown('<div class="warn-box">⚠️ No citations found. Please upload reference PDFs or modify your topic.</div>',
                    unsafe_allow_html=True)

    st.markdown("**Search for more citations:**")
    col_f1, col_f2 = st.columns([3,1])
    fetch_query = col_f1.text_input("", placeholder="e.g. working memory children India",
                                     key="fetch_q", label_visibility="collapsed")
    free_left = max(0, 2 - st.session_state.get("fetch_more_count", 0))
    btn_label = f"Fetch ({'Free' if free_left > 0 else '-5 credits'})"
    if col_f2.button(btn_label, key="do_fetch"):
        q = st.session_state.get("fetch_q", "").strip()
        if not q:
            st.warning("Enter a search term.")
        else:
            with st.spinner(f"Searching for: {q}"):
                try:
                    new_papers = fetch_citation_bank([q], target=8)
                    if len(new_papers) < 2:
                        broad = " ".join(q.split()[:3])
                        new_papers += fetch_citation_bank([broad], target=5)
                    existing = {p.get("title","")[:40].lower() for p in st.session_state.citation_bank}
                    added = [p for p in new_papers if p.get("title","")[:40].lower() not in existing]
                    st.session_state.citation_bank.extend(added)
                    st.session_state["fetch_more_count"] = st.session_state.get("fetch_more_count", 0) + 1
                    if added:
                        st.success(f"Added {len(added)} citations.")
                    else:
                        st.warning("No new citations found. Try different terms.")
                    st.rerun()
                except Exception as ex:
                    st.error(f"Fetch error: {ex}")


    # Conference theme matching
    if st.session_state.conference_rules and EXTENSIONS_AVAILABLE:
        st.markdown("### 🎯 Conference Theme Matching")
        if not st.session_state.conference_themes:
            with st.spinner("Matching your domain to conference themes..."):
                ct = extract_conference_themes(st.session_state.conference_rules)
                st.session_state.conference_themes = ct
                if ct.get("themes"):
                    matches = match_domain_to_themes(
                        da.get("detected_domain",""), st.session_state.topic,
                        ct["themes"]
                    )
                    st.session_state.conference_match = matches

        ct = st.session_state.conference_themes
        if ct.get("conference_name"):
            st.markdown(f"**Conference:** {ct['conference_name']}")
        if ct.get("word_limit"):
            st.markdown(f'<div class="warn-box">📏 Conference word limit: <strong>{ct["word_limit"]:,}</strong> words</div>', unsafe_allow_html=True)

        matches = st.session_state.conference_match
        if matches:
            st.markdown("**Best matching themes (AI ranked):**")
            for i, m in enumerate(matches[:3]):
                conf_pct = m.get("confidence", 0)
                bar = "█" * int(conf_pct/10) + "░" * (10-int(conf_pct/10))
                st.markdown(f"{i+1}. **{m['theme']}** — `{bar}` {conf_pct:.0f}% match")
            
            # Title suggestions
            if matches:
                title_suggestions = suggest_paper_title(
                    da.get("detected_domain",""), st.session_state.topic,
                    matches[0]["theme"],
                    [h.get("alternate","") for h in st.session_state.hypotheses[:2]]
                )
                st.markdown("**Suggested titles aligned with conference theme:**")
                for ts in title_suggestions:
                    if st.button(f"Use: {ts[:80]}", key=f"title_{ts[:20]}"):
                        st.session_state.topic = ts
                        st.rerun()
                st.text_input("Or type your own title:", key="custom_conf_title",
                               placeholder="Your custom title here")

    # Fetch more citations
    st.markdown("---")
    _fetch_input = st.text_input("Search more citations", 
                                  placeholder="e.g. AI tutoring working memory India",
                                  key="fetch_extra_input")
    _fetch_count = st.session_state.get("fetch_more_count", 0)
    _fetch_cost  = 0 if _fetch_count < 2 else 5
    _free_left   = max(0, 2 - _fetch_count)
    _fetch_label = f"+ Fetch More ({'Free — '+str(_free_left)+' left' if _fetch_cost == 0 else f'-{_fetch_cost} credits'})"
    if st.button(_fetch_label, use_container_width=True, key="fetch_main_btn"):
        _term = st.session_state.get("fetch_extra_input","").strip()
        if not _term:
            st.warning("Enter a search term above.")
        else:
            with st.spinner(f"Fetching citations for: {_term}"):
                from citation_engine import fetch_citation_bank
                _more = fetch_citation_bank([_term], target=5)
                if len(_more) < 2:
                    _broad = " ".join(_term.split()[:3])
                    _more2 = fetch_citation_bank([_broad], target=5)
                    _more.extend(_more2)
                _exist = {p.get("title","")[:40].lower() for p in st.session_state.citation_bank}
                _new   = [p for p in _more if p.get("title","")[:40].lower() not in _exist]
                st.session_state.citation_bank.extend(_new)
                st.session_state.fetch_more_count = _fetch_count + 1
                if _new:
                    st.success(f"✅ Added {len(_new)} new citations.")
                else:
                    st.warning("No new citations found. Try different terms.")
                st.rerun()
    st.markdown("---")

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=2; st.rerun()
    with c2:
        if st.button("▶ Set Objectives & Hypotheses →", use_container_width=True):
            st.session_state.step=4; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 4 — OBJECTIVES & HYPOTHESES
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 4:
    st.markdown("## Step 4 — Objectives & Hypotheses")
    st.markdown("AI suggests 3+3. Edit, refresh, add, or prompt your own. These drive the ENTIRE paper structure.")

    da = st.session_state.domain_analysis
    cites_text = bank_to_prompt_text(st.session_state.citation_bank[:6], st.session_state.citation_style)

    def gen_obj_hyp(custom_prompt: str = ""):
        with st.spinner("Generating objectives and hypotheses..."):
            prompt = f"""Topic: "{st.session_state.topic}"
Domain: {da.get("detected_domain","")}
Citations available: {cites_text[:1500]}
{f"Custom instruction: {custom_prompt}" if custom_prompt else ""}

Generate 3 research objectives and 3 testable hypotheses.
Return ONLY valid JSON:
{{
  "objectives": [
    {{"id": "O1", "text": "To examine...", "basis": "brief rationale"}},
    {{"id": "O2", "text": "To investigate...", "basis": "..."}},
    {{"id": "O3", "text": "To assess...", "basis": "..."}}
  ],
  "hypotheses": [
    {{"id": "H1", "null": "There is no significant...", "alternate": "There is a significant...", "supported_by": "citation or theory"}},
    {{"id": "H2", "null": "...", "alternate": "...", "supported_by": "..."}},
    {{"id": "H3", "null": "...", "alternate": "...", "supported_by": "..."}}
  ]
}}
Make them domain-specific and testable. No generic placeholders."""
            try:
                resp = claude_call(MASTER_SYSTEM, prompt, 2000, cheap=False)
                data = json_parse(resp)
                st.session_state.objectives = data.get("objectives", [])
                st.session_state.hypotheses = data.get("hypotheses", [])
            except Exception as e:
                st.error(f"Generation error: {e}")

    if not st.session_state.objectives:
        gen_obj_hyp()
        st.rerun()

    # Display and edit objectives
    st.markdown("### Objectives")
    for i, obj in enumerate(st.session_state.objectives):
        c1,c2,c3 = st.columns([4,1,1])
        new_text = c1.text_input(f"O{i+1}", value=obj.get("text",""), key=f"obj_{i}", label_visibility="collapsed")
        st.session_state.objectives[i]["text"] = new_text
        if c2.button("🔄", key=f"robj_{i}"):
            gen_obj_hyp(f"Regenerate objective {i+1} only, keep others similar")
            st.rerun()
        if c3.button("✕", key=f"dobj_{i}"):
            st.session_state.objectives.pop(i); st.rerun()

    if st.button("+ Add Objective"):
        st.session_state.objectives.append({"id":f"O{len(st.session_state.objectives)+1}", "text":"To explore...", "basis":""})
        st.rerun()

    # Display and edit hypotheses
    st.markdown("### Hypotheses")
    for i, hyp in enumerate(st.session_state.hypotheses):
        with st.expander(f"H{i+1}: {hyp.get('alternate','')[:80]}..."):
            alt = st.text_area("Alternate hypothesis", value=hyp.get("alternate",""), key=f"hyp_{i}", height=60)
            null = st.text_input("Null hypothesis", value=hyp.get("null",""), key=f"hyp_null_{i}")
            st.session_state.hypotheses[i]["alternate"] = alt
            st.session_state.hypotheses[i]["null"] = null
            c1,c2 = st.columns(2)
            if c1.button("🔄 Refresh this hypothesis", key=f"rhyp_{i}"):
                gen_obj_hyp(f"Regenerate H{i+1} only. Make it different from: {hyp.get('alternate','')}")
                st.rerun()
            if c2.button("✕ Delete", key=f"dhyp_{i}"):
                st.session_state.hypotheses.pop(i); st.rerun()

    if st.button("+ Add Hypothesis"):
        st.session_state.hypotheses.append({"id":f"H{len(st.session_state.hypotheses)+1}",
                                              "null":"There is no significant...",
                                              "alternate":"There is a significant...",
                                              "supported_by":""})
        st.rerun()

    # Custom prompt to modify all
    custom = st.text_input("Prompt to modify all (e.g. 'Focus on gender moderation', 'Add longitudinal perspective')")
    c1,c2 = st.columns(2)
    with c1:
        if st.button("🔄 Master Refresh (All)", use_container_width=True):
            gen_obj_hyp(custom)
            st.rerun()
    with c2:
        if st.button("▶ Build Structure →", use_container_width=True):
            st.session_state.step=5; st.rerun()

    if st.button("◀ Back"): st.session_state.step=3; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 5 — STRUCTURE
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 5:
    st.markdown("## Step 5 — Paper Structure")
    da = st.session_state.domain_analysis
    tmpl = bank.templates.get(st.session_state.template_key, {})
    mandatory = bank.get_mandatory_sections(st.session_state.template_key)
    conf_rules = st.session_state.conference_constraints

    if not st.session_state.structure:
        with st.spinner("Building structure..."):
            hyp_text = json.dumps(st.session_state.hypotheses, indent=2)
            obj_text = json.dumps(st.session_state.objectives, indent=2)
            prompt = f"""Build a paper structure for:
Topic: {st.session_state.topic}
Document type: {st.session_state.paper_type}
Domain: {da.get("detected_domain","")}
Word limit: {st.session_state.word_limit}
Objectives: {obj_text}
Hypotheses: {hyp_text}
Mandatory sections (must include): {mandatory}
Conference constraints: {json.dumps(conf_rules)}
User overrides: {st.session_state.user_overrides or 'None'}

Return ONLY a JSON array:
[
  {{"section": "Abstract", "word_allocation": 250, "notes": "structured: Background/Methods/Results/Conclusion"}},
  ...
]
Total must not exceed {st.session_state.word_limit}. Every section must trace to at least one objective or hypothesis.
Return ONLY valid JSON array. No markdown fences."""
            try:
                resp = claude_call(MASTER_SYSTEM, prompt, 1500)
                st.session_state.structure = json_parse(resp)
            except Exception as e:
                st.error(f"Structure error: {e}")

    if st.session_state.structure:
        total = sum(s.get("word_allocation",0) for s in st.session_state.structure)
        st.markdown(f"**Total: `{total:,}` / `{st.session_state.word_limit:,}` words**")

        edited = []
        for i, sec in enumerate(st.session_state.structure):
            c1,c2,c3 = st.columns([3,1,1])
            w = c1.number_input(sec.get("section",""), value=max(50, int(sec.get("word_allocation") or 500)),
                                min_value=50, max_value=4000, key=f"sec_{i}")
            c2.caption(sec.get("notes","")[:60])
            del_sec = c3.button("✕", key=f"dsec_{i}")
            if not del_sec:
                edited.append({**sec, "word_allocation": w})

        if st.button("+ Add Section"):
            new_sec = st.text_input("New section name", key="new_sec_name")
            if new_sec:
                edited.append({"section": new_sec, "word_allocation": 500, "notes": ""})

        st.session_state.structure = edited

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=4; st.rerun()
    with c2:
        if st.button("▶ Statistics Plan →", use_container_width=True):
            st.session_state.step=6; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 6 — STATISTICS PLAN
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 6:
    st.markdown("## Step 6 — Statistics Plan")
    da = st.session_state.domain_analysis

    if not st.session_state.stats_plan:
        with st.spinner("Building statistics plan for your domain..."):
            hyp_text = json.dumps([h.get("alternate","") for h in st.session_state.hypotheses], indent=2)
            prompt = f"""Topic: {st.session_state.topic}
Domain: {da.get("detected_domain","")}
Hypotheses: {hyp_text}
Variables: {da.get("key_variables",[])}
Uploaded data context: {st.session_state.uploaded_material[:1000] or 'None'}
User overrides: {st.session_state.user_overrides or 'None'}

Return ONLY this JSON, no other text:
{{"Basic":{{"tests":["t-test","chi-square"],"rationale":"appropriate for domain","limitations":"limited depth","journal_acceptance_probability":"60-70%","recommended_diagrams":["bar chart"]}},
"Medium":{{"tests":["ANOVA","correlation","Cronbach alpha"],"rationale":"standard academic","limitations":"no SEM","journal_acceptance_probability":"75-85%","recommended_diagrams":["violin plot","heatmap"]}},
"Advanced":{{"tests":["SEM","CFA","Bonferroni","bootstrap CI"],"rationale":"publication grade","limitations":"none significant","journal_acceptance_probability":"88-95%","recommended_diagrams":["SEM path diagram","radar chart"]}},
"ai_recommendation":"Medium",
"recommendation_reason":"Best fit for this domain and sample size"}}"""
            try:
                resp = claude_call(MASTER_SYSTEM, prompt, 1500)
                st.session_state.stats_plan = json_parse(resp)
            except Exception as e:
                st.error(f"Stats plan error: {e}")

    sp = st.session_state.stats_plan
    if sp and "Basic" in sp:
        rec = sp.get("ai_recommendation","Medium")
        st.markdown(f'<div class="ok-box">🤖 <strong>AI Recommends: {rec}</strong> — {sp.get("recommendation_reason","")}</div>',
                    unsafe_allow_html=True)

        cols = st.columns(3)
        for col, level, emoji in zip(cols, ["Basic","Medium","Advanced"], ["💰","⚖️","🏆"]):
            with col:
                data = sp.get(level, {})
                st.markdown(f"### {emoji} {level}")
                st.markdown(f"**Acceptance:** `{data.get('journal_acceptance_probability','?')}`")
                for t in data.get("tests", []):
                    st.markdown(f"- {t}")
                if data.get("recommended_diagrams"):
                    st.markdown("**Diagrams:**")
                    for d in data["recommended_diagrams"]:
                        st.markdown(f"- {d}")
                if data.get("limitations"):
                    st.markdown(f'<div class="warn-box">⚠️ {data["limitations"]}</div>', unsafe_allow_html=True)

        choice = st.radio("Select level", ["Basic","Medium","Advanced"],
                           index=["Basic","Medium","Advanced"].index(rec), horizontal=True)
        st.session_state.stats_level = choice

        if choice == "Advanced":
            st.info("📊 Diagrams included by default at Advanced.")
            st.session_state.include_diagrams = not st.toggle("Disable diagrams", value=False)
        else:
            st.session_state.include_diagrams = st.toggle("Include diagrams?", value=False)

    # ── Questionnaire / Likert Construct Editor (optional add-on) ──────
    if EXTENSIONS_AVAILABLE:
        st.markdown("---")
        st.markdown("### 📋 Questionnaire / Likert Constructs *(optional)*")
        st.caption(
            "Generate AI-drafted Likert items per construct. "
            "Minimum 5 items per construct enforced. Indian standard: 8+. "
            "Named CSV columns (e.g. WMP_Q1, WMP_Q2) are generated from construct names.")

        col_gen, col_info = st.columns([2, 1])
        with col_info:
            st.markdown(
                '<div class="ok-box" style="font-size:12px">'
                '✅ Free for Medium+ tiers<br>'
                '⚠️ Basic: −16 credits (₹80)</div>',
                unsafe_allow_html=True)

        with col_gen:
            n_constructs = st.number_input(
                "Number of constructs", min_value=1, max_value=6,
                value=len(st.session_state.hypotheses) or 3,
                key="n_constructs_input",
                help="Usually one per hypothesis.")
            items_per = st.number_input(
                "Items per construct", min_value=5, max_value=20, value=8,
                key="items_per_construct",
                help="Minimum 5. Indian studies standard: 8.")

        if st.button("🤖 Generate Constructs", key="gen_constructs_btn",
                     use_container_width=True):
            with st.spinner("Generating Likert constructs…"):
                hyp_texts = [h.get("alternate","") for h in st.session_state.hypotheses]
                domain_str = st.session_state.domain_analysis.get("detected_domain","Research")
                prompt = generate_constructs_prompt(
                    domain_str, hyp_texts,
                    n_constructs=int(n_constructs),
                    items_per_construct=int(items_per))
                try:
                    raw = claude_call(MASTER_SYSTEM, prompt, 3000, cheap=False)
                    raw = re.sub(r'```json|```', '', raw).strip()
                    data = json.loads(raw)
                    st.session_state.constructs = data.get("constructs", [])
                    st.session_state.demo_items  = data.get("demographic_items", [])
                    st.success(
                        f"✅ {len(st.session_state.constructs)} constructs generated. "
                        f"Edit below before proceeding.")
                    st.rerun()
                except Exception as _e:
                    st.error(f"Construct generation error: {_e}")

        if st.session_state.get("constructs"):
            st.session_state.constructs = render_construct_editor(
                st.session_state.constructs,
                claude_call_fn=claude_call,
                domain=st.session_state.domain_analysis.get("detected_domain",""),
                hypotheses=[h.get("alternate","") for h in st.session_state.hypotheses],
            )
            st.caption(
                f"✅ {len(st.session_state.constructs)} constructs ready. "
                "CSV columns will be named from construct IDs (e.g. C1Q1, C1Q2…).")
        else:
            st.caption("No constructs yet. Generate above, or skip to proceed without questionnaire.")

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=5; st.rerun()
    with c2:
        if st.button("▶ Generate Narratives →", use_container_width=True):
            st.session_state.step=7; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 7 — NARRATIVES (Last free step)
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 7:
    st.markdown("## Step 7 — Choose Your Story")
    st.markdown("Three possible outcomes. Pick the one that fits your purpose. **This is the last free step.**")

    da = st.session_state.domain_analysis
    hyp_list = [h.get("alternate","") for h in st.session_state.hypotheses]

    if not st.session_state.narratives:
        with st.spinner("Generating three narrative outcomes..."):
            st.session_state.narratives = generate_narrative_targets(
                da.get("detected_domain","Generic"), hyp_list, st.session_state.stats_level)

    narratives = st.session_state.narratives
    selected = st.session_state.selected_narrative

    for key in ["A","B","C"]:
        n = narratives.get(key, {})
        is_sel = selected == key
        card_class = "narrative-card narrative-selected" if is_sel else "narrative-card"

        st.markdown(f'<div class="{card_class}">', unsafe_allow_html=True)
        st.markdown(f"### {'★ ' if is_sel else ''}Narrative {key}: {n.get('label','')}")
        st.markdown(n.get("description",""))

        st.markdown("**Hypothesis outcomes:**")
        for ht in n.get("hypotheses",[]):
            icon = "✅" if ht.get("supported") else "❌"
            st.markdown(f"{icon} {ht.get('hypothesis','')[:70]}... — **{ht.get('effect_label','')} effect**")

        if st.button(f"Select Narrative {key}", key=f"sel_{key}", use_container_width=True):
            st.session_state.selected_narrative = key
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    custom_prompt = st.text_input("Want a different narrative? Describe it:")
    if st.button("🔄 Generate Custom Narrative"):
        with st.spinner("Generating..."):
            st.session_state.narratives = generate_narrative_targets(
                da.get("detected_domain","Generic"), hyp_list, st.session_state.stats_level)
        st.rerun()

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=6; st.rerun()
    with c2:
        if st.button("▶ Generate Data & Paper →", use_container_width=True,
                     type="primary"):
            if not st.session_state.selected_narrative:
                st.warning("Select a narrative first.")
            else:
                # ── CREDIT DEDUCTION HERE ──────────────────
                # Credit costs: 1 credit = ₹10. Paper costs by tier.
                cost = {"Basic":104,"Medium":160,"Advanced":240,"Premium":400,"Ultra":600}.get(st.session_state.tier, 104)
                st.session_state["live_balance"] = max(0, st.session_state.get("live_balance", 9999.0) - cost)
                show_credit_drain(cost, f"{st.session_state.tier} paper — {cost} credits")
                st.session_state.step=8
                st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 8 — REVERSE-ENGINEER DATA
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 8:
    st.markdown("## Step 8 — Generating Dataset")

    narrative_key = st.session_state.selected_narrative
    narrative = st.session_state.narratives.get(narrative_key, {})

    if st.session_state.synthetic_df is None:
        with st.spinner("Reverse-engineering dataset to match target statistics..."):
            try:
                # Inject user's sample size into narrative
                _n = int(st.session_state.get("sample_size_override") or 90)
                narrative["n"] = _n
                df = reverse_engineer_dataset(
                    narrative,
                    constructs=st.session_state.get("constructs") or None
                )
                verification = verify_statistics(df, narrative)
                st.session_state.synthetic_df = df
                st.session_state.stats_verification = verification
                st.session_state.target_stats = narrative
            except Exception as e:
                st.error(f"Data generation error: {e}")
                credits.refund(st.session_state.user_id,
                                CREDIT_COSTS.get(st.session_state.tier,1.0),
                                st.session_state.paper_id)
                st.stop()

    df = st.session_state.synthetic_df
    v  = st.session_state.stats_verification

    st.markdown(f"**Dataset: {len(df)} rows × {len(df.columns)} columns**")

    # Show verification report
    st.markdown("### Statistics Verification")
    ca = v.get("cronbach_alpha",{})
    if ca:
        icon = "✅" if ca.get("pass") else "⚠️"
        st.markdown(f"{icon} Cronbach's α: computed `{ca.get('computed','-')}` | target `{ca.get('target','-')}`")

    for hv in v.get("hypotheses",[]):
        p_icon = "✅" if hv.get("p_pass") else "⚠️"
        d_icon = "✅" if hv.get("d_pass") else "⚠️"
        st.markdown(f"H{hv['hypothesis_num']}: t={hv.get('t_statistic','-')}, "
                    f"p={hv.get('p_value','-')} {p_icon}, "
                    f"d={hv.get('cohens_d','-')} {d_icon}")

    with st.expander("Preview dataset (first 10 rows)"):
        st.dataframe(df.head(10))

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=7; st.rerun()
    with c2:
        if st.button("▶ Write Paper →", use_container_width=True):
            st.session_state.step=9; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 9 — FULL PAPER GENERATION
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 9:
    st.markdown("## Step 9 — Writing Paper")

    if st.session_state.full_paper:
        st.markdown('<div class="ok-box">Paper already generated. Proceed to review or regenerate below.</div>',
                    unsafe_allow_html=True)
    else:
        import random as _random
        quote = _random.choice(STATISTICIAN_QUOTES)
        st.markdown(f'''<div class="quote-box">📖 {quote}</div>''', unsafe_allow_html=True)
        progress_bar = st.progress(0)
        progress_msg = st.empty()
        with st.spinner(""):
            da        = st.session_state.domain_analysis
            narrative = st.session_state.narratives.get(st.session_state.selected_narrative, {})
            v         = st.session_state.stats_verification
            structure = st.session_state.structure
            cites_text   = bank_to_prompt_text(st.session_state.citation_bank[:12],
                                               st.session_state.citation_style)
            stat_results = _build_stats_text(v, narrative)

            try:
                # ── SECTIONAL GENERATION (fixes truncation at ~1,400w) ──────
                # Each section is its own API call with a focused token budget.
                # Concatenated result reliably reaches target word count on
                # any model tier (Haiku → Sonnet → Opus).
                paper = generate_paper_sectional(
                    topic                 = st.session_state.topic,
                    paper_type            = st.session_state.paper_type,
                    tier                  = st.session_state.tier,
                    domain                = da.get("detected_domain", ""),
                    language              = st.session_state.language,
                    grade                 = st.session_state.language_complexity,
                    tone                  = da.get("tone_recommendation", "formal"),
                    structure             = structure,
                    stat_results          = stat_results,
                    objectives            = st.session_state.objectives,
                    hypotheses            = st.session_state.hypotheses,
                    cites_text            = cites_text,
                    uploaded_material     = st.session_state.uploaded_material,
                    conference_constraints= st.session_state.conference_constraints,
                    user_overrides        = st.session_state.user_overrides,
                    citation_style        = st.session_state.citation_style,
                    stats_level           = st.session_state.stats_level,
                    include_diagrams      = st.session_state.include_diagrams,
                    citation_bank         = st.session_state.citation_bank,
                    progress_bar          = progress_bar,
                    progress_msg          = progress_msg,
                )

                # ── Phase B: citation discipline ─────────────────────────────
                progress_msg.markdown("⚙️ Enforcing citation discipline…")
                progress_bar.progress(90)
                paper, removed = enforce_citation_discipline(paper, st.session_state.citation_bank)

                # ── Phase C: clean markdown artefacts ───────────────────────
                if FORMATTER_AVAILABLE:
                    paper = clean_markdown(paper)

                progress_bar.progress(100)
                wc_final = word_count(paper)
                progress_msg.markdown(
                    f"✅ Paper complete — **{wc_final:,} words** written "
                    f"(target: {st.session_state.word_limit:,})"
                )
                if removed:
                    st.info(f"ℹ️ {removed} unverified citation(s) automatically removed.")

                # ── warn if still short (should not happen post-fix) ─────────
                if wc_final < int(st.session_state.word_limit * 0.80):
                    st.warning(
                        f"⚠️ Paper generated {wc_final:,} words against a "
                        f"{st.session_state.word_limit:,}-word target. "
                        "Use 'Regenerate Section' below to pad any short sections."
                    )

                st.session_state.full_paper = paper
                credits.confirm(st.session_state.user_id, st.session_state.paper_id)

            except Exception as e:
                st.error(f"Generation error: {e}")
                credits.refund(st.session_state.user_id,
                               CREDIT_COSTS.get(st.session_state.tier, 1.0),
                               st.session_state.paper_id)
                st.stop()

    wc = word_count(st.session_state.full_paper)
    st.markdown(f"**Generated: `{wc:,}` words** (budget: `{st.session_state.word_limit:,}`)")

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=8; st.rerun()
    with c2:
        if st.button("▶ Quality Audit →", use_container_width=True):
            st.session_state.step=10; st.rerun()




# ══════════════════════════════════════════════════════════
# STEP 10 — AUDIT
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 10:
    st.markdown("## Step 10 — Quality Audit")
    st.markdown("Fireworks gpt-oss-20b → Groq Llama-70b → Haiku fallback. Claude does not audit its own output.")

    if not st.session_state.audit_issues:
        with st.spinner("Running data integrity + neutrality audit..."):
            api_keys = get_api_keys()
            # Audit uses Fireworks→Groq→Haiku — never Sonnet auditing itself
            cleaned, issues, passed, rewrite_brief = audit_and_clean(
                st.session_state.full_paper, api_keys)
            st.session_state.full_paper = cleaned
            st.session_state.audit_issues = issues

            if not passed and rewrite_brief:
                st.warning("Audit found issues. Auto-rewriting flagged sections...")
                with st.spinner("Rewriting..."):
                    rewrite_prompt = f"{rewrite_brief}\n\nREWRITE THE FLAGGED SECTIONS ONLY:\n{st.session_state.full_paper[:6000]}"
                    try:
                        fixed = claude_call(MASTER_SYSTEM, rewrite_prompt, 3000)
                        st.session_state.full_paper = fixed
                        st.session_state.audit_issues.append("Auto-rewrite applied.")
                    except Exception as e:
                        st.error(f"Rewrite error: {e}")

    if st.session_state.audit_issues:
        st.markdown("**Issues processed:**")
        for issue in st.session_state.audit_issues[:10]:
            st.caption(f"• {issue}")
    else:
        st.markdown('<div class="ok-box">✅ Audit passed. No issues found.</div>', unsafe_allow_html=True)

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=9; st.rerun()
    with c2:
        if st.button("▶ Review & Edit →", use_container_width=True):
            st.session_state.step=11; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 11 — HUMAN REVIEW & SECTION REGENERATION
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 11:
    st.markdown("## Step 11 — Review & Edit")

    wc = word_count(st.session_state.full_paper)
    can_regen, remaining = credits.check_max_regens(
        st.session_state.user_id, st.session_state.paper_id, st.session_state.tier)

    st.markdown(f"**Words: `{wc:,}`** | Section regenerations remaining: **{remaining}**")
    regen_cost = 16
    st.markdown(f"*First 2 regenerations free · After that: −{regen_cost} credits (₹{regen_cost*5}) each*")

    with st.expander("📄 Full Paper Preview", expanded=True):
        # Show clean preview — no markdown artifacts
        preview_text = st.session_state.full_paper
        if FORMATTER_AVAILABLE:
            preview_text = clean_markdown(preview_text)
        # Render as clean text with section breaks
        sections = preview_text.split("\n\n")
        for sec in sections:
            sec = sec.strip()
            if not sec:
                continue
            # Detect headings
            if len(sec.split()) <= 8 and not sec.endswith(".") and sec[0].isupper():
                st.markdown(f"**{sec}**")
            else:
                st.write(sec)

    st.markdown("### Regenerate a Section")
    section_names = [s.get("section","") for s in st.session_state.structure]
    target_section = st.selectbox("Which section to regenerate?", section_names)
    regen_prompt = st.text_area("What to change?", height=70,
                                 placeholder="e.g. 'Make discussion more critical', 'Add more citations to intro', 'Expand methodology detail'")

    if st.button("🔄 Regenerate Section"):
        if not can_regen:
            st.error("Maximum regenerations reached for this paper.")
        else:
            allowed, cost, msg = credits.use_regen(
                st.session_state.user_id, st.session_state.paper_id)
            if not allowed:
                st.error(msg)
            else:
                if cost > 0:
                    st.session_state["live_balance"] = max(0, st.session_state.get("live_balance", 10.0) - cost)
                    show_credit_drain(cost, f"section regen")
                with st.spinner(f"Rewriting {target_section}..."):
                    rp = (f"Rewrite the '{target_section}' section of this paper.\n"
                          f"Instruction: {regen_prompt}\n"
                          f"Keep all other sections intact.\n"
                          f"Use ONLY these citations: {bank_to_prompt_text(st.session_state.citation_bank[:8], st.session_state.citation_style)}\n\n"
                          f"ORIGINAL PAPER:\n{st.session_state.full_paper[:6000]}")
                    try:
                        new_section = claude_call(MASTER_SYSTEM, rp, 2000)
                        new_section, _ = enforce_citation_discipline(
                            new_section, st.session_state.citation_bank)
                        st.session_state.full_paper = _replace_section(
                            st.session_state.full_paper, target_section, new_section)
                        st.success(f"✅ {target_section} rewritten. {msg}")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Regen error: {e}")

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=10; st.rerun()
    with c2:
        if st.button("▶ Sign MSATA & Download →", use_container_width=True):
            st.session_state.step=12; st.rerun()




# ══════════════════════════════════════════════════════════
# STEP 12 — MSATA SIGN
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 12:
    if MSATA_AVAILABLE:
        signed = render_msata_step(paper_title=st.session_state.topic)
        if signed:
                credits.log_msata(
                    st.session_state.user_id, st.session_state.paper_id,
                    st.session_state.topic, "IP_TBD", "Browser"
                )
                st.session_state.msata_signed = True
                st.session_state.step = 13
                st.rerun()
    else:
        # Fallback simple MSATA
        st.markdown("## Step 12 — Agreement & Download")
        st.markdown('<div class="warn-box"><strong>MASTER SERVICE & ASSET TRANSFER AGREEMENT</strong><br>By proceeding you acknowledge this is a Technical Prototype (UGC 2018). You take full responsibility.</div>', unsafe_allow_html=True)
        user_name  = st.text_input("Full Name")
        user_email = st.text_input("Email")
        if st.button("✅ I Accept — Unlock Download", use_container_width=True):
            if user_name and user_email:
                st.session_state.msata_signed = True
                st.session_state.step = 13
                st.rerun()
    if st.button("◀ Back"): st.session_state.step=11; st.rerun()

# ══════════════════════════════════════════════════════════
# STEP 13 — DOWNLOAD
# ══════════════════════════════════════════════════════════
elif st.session_state.step == 13:
    st.markdown("## Step 13 — Download")

    # Table and diagram options
    col_t1, col_t2, col_t3 = st.columns(3)
    with col_t1:
        table_mode = st.radio("Tables", ["Combined", "Separate (SPSS-style)"], horizontal=True)
        st.session_state.table_mode = table_mode.lower().split()[0]
    with col_t2:
        show_diag = st.toggle("Include Diagrams", value=st.session_state.show_diagrams)
        st.session_state.show_diagrams = show_diag
    with col_t3:
        if show_diag and st.button("📊 Preview Diagrams"):
            df = st.session_state.synthetic_df
            if df is not None and EXTENSIONS_AVAILABLE:
                try:
                    from diagram_engine import generate_figures_for_paper
                    figs = generate_figures_for_paper(
                        df=df,
                        stats_verification=st.session_state.get("stats_verification",{}),
                        narrative=st.session_state.narratives.get(
                            st.session_state.selected_narrative or "A", {}),
                        domain=st.session_state.domain_analysis.get("detected_domain","")
                            if isinstance(st.session_state.domain_analysis,dict) else ""
                    )
                    for fig_bytes, fig_cap in figs[:3]:
                        st.image(fig_bytes, caption=fig_cap)
                except Exception as _fe:
                    st.warning(f"Preview error: {_fe}")

    if not st.session_state.msata_signed:
        st.warning("Please sign the MSATA first.")
        st.session_state.step=12; st.rerun()

    # Phase C: final citation verification
    with st.spinner("Running final citation verification (Phase C)..."):
        verified, stripped = verify_bank_phase_c(st.session_state.citation_bank)
        if stripped:
            st.info(f"ℹ️ {stripped} citation(s) could not be CrossRef-verified and were reviewed.")
        st.session_state.citation_bank = verified

    st.markdown('<div class="ok-box">✅ Paper ready. Download below.</div>', unsafe_allow_html=True)

    paper = st.session_state.full_paper
    slug  = re.sub(r'[^a-z0-9]+','_', st.session_state.topic.lower())[:40]

    fmt_style = st.session_state.get("formatting_style", "SPPU")
    kw = extract_keywords(paper) if FORMATTER_AVAILABLE else []
    
    # Build stats table for DOCX
    v2 = st.session_state.get("stats_verification", {})
    narrative2 = st.session_state.narratives.get(st.session_state.selected_narrative, {})
    stats_tbl_headers = ["Hypothesis", "Test", "t / F", "df", "p-value", "Effect Size", "Decision"]
    stats_tbl_rows = []
    for hv in v2.get("hypotheses", []):
        idx = hv["hypothesis_num"] - 1
        ht2 = (narrative2.get("hypotheses") or [{}]*3)
        h2 = ht2[idx] if idx < len(ht2) else {}
        stats_tbl_rows.append([
            f"H{hv['hypothesis_num']}",
            "Paired t-test",
            str(hv.get("t_statistic", "-")),
            str(narrative2.get("n", 90) - 1),
            str(hv.get("p_value", "-")),
            f"d = {hv.get('cohens_d', '-')} ({h2.get('effect_label', '')})",
            "Supported" if h2.get("supported") else "Not Supported",
        ])
    ca2 = v2.get("cronbach_alpha", {})
    if ca2:
        stats_tbl_rows.append([
            "Scale Reliability", "Cronbach Alpha",
            str(ca2.get("computed", "-")), "-", "-",
            "Good" if (ca2.get("computed") or 0) >= 0.79 else "Acceptable",
            "Pass" if ca2.get("pass") else "Review",
        ])
    if stats_tbl_rows:
        stats_tbl_rows.append(["Total", f"{len(v2.get('hypotheses',[]))} hypotheses",
                                 "-", "-", "-", "-",
                                 f"{sum(1 for r in stats_tbl_rows[:-0] if r[-1]=='Supported')}/{len(v2.get('hypotheses',[]))} Supported"])
    
    # Figures from session
    figures_list = st.session_state.get("docx_figures", [])
    
    # Generate diagrams — chart types decided dynamically by data
    if st.session_state.get("show_diagrams") and EXTENSIONS_AVAILABLE and st.session_state.synthetic_df is not None:
        try:
            auto_figures = generate_figures_for_paper(
                df=st.session_state.synthetic_df,
                stats_verification=st.session_state.get("stats_verification", {}),
                narrative=st.session_state.narratives.get(
                    st.session_state.selected_narrative, {}),
                domain=st.session_state.domain_analysis.get("detected_domain","")
                    if isinstance(st.session_state.domain_analysis, dict) else "",
            )
            figures_list.extend(auto_figures)
        except Exception as _fig_err:
            pass  # Never block download due to figure error
    
    stats_tables_for_docx = [(stats_tbl_headers, stats_tbl_rows, "Table 1. Summary of Hypothesis Testing Results")] if stats_tbl_rows else None
    
    if FORMATTER_AVAILABLE:
        docx_bytes = build_professional_docx(
            content=paper,
            title=st.session_state.topic,
            style_key=fmt_style,
            keywords=kw,
            citation_style=st.session_state.citation_style,
            figures=figures_list if figures_list else None,
            stats_tables=stats_tables_for_docx,
        )
    else:
        docx_bytes = build_docx(paper, st.session_state.topic)
    # CSV — raw dataset
    df_raw = st.session_state.synthetic_df
    csv_bytes = df_raw.to_csv(index=False).encode("utf-8") if df_raw is not None else b""
    
    # Stats verification CSV — optional, costs 0.50 credits
    st.markdown("---")
    st.markdown("### 📊 Verified Statistics CSV")
    st.markdown("Download a CSV confirming all paper statistics match the dataset. Costs **0.50 credits**.")
    
    if st.button("📥 Generate Stats Verification CSV (−0.50 credits)"):
        stat_cost = 10
        st.session_state["live_balance"] = max(0, st.session_state.get("live_balance", 10.0) - stat_cost)
        show_credit_drain(stat_cost, "stats verification")
        
        v = st.session_state.get("stats_verification", {})
        narrative = st.session_state.narratives.get(st.session_state.selected_narrative, {})
        
        import pandas as _pd
        rows = []
        for hv in v.get("hypotheses", []):
            idx = hv["hypothesis_num"] - 1
            ht = (narrative.get("hypotheses") or [{}]*3)
            h = ht[idx] if idx < len(ht) else {}
            rows.append({
                "Hypothesis": f"H{hv['hypothesis_num']}",
                "Test": "Paired t-test",
                "t_statistic": hv.get("t_statistic"),
                "df": narrative.get("n", 90) - 1,
                "p_value": hv.get("p_value"),
                "Cohens_d": hv.get("cohens_d"),
                "Effect_Size": h.get("effect_label", ""),
                "Decision": "Supported" if h.get("supported") else "Not Supported",
                "Target_d": h.get("cohens_d"),
                "Match": "✓" if hv.get("d_pass") else "⚠",
            })
        
        ca = v.get("cronbach_alpha", {})
        rows.append({
            "Hypothesis": "Scale Reliability",
            "Test": "Cronbach Alpha",
            "t_statistic": ca.get("computed"),
            "df": "",
            "p_value": "",
            "Cohens_d": "",
            "Effect_Size": "Good" if (ca.get("computed") or 0) >= 0.79 else "Acceptable",
            "Decision": "Acceptable" if (ca.get("computed") or 0) >= 0.70 else "Poor",
            "Target_d": ca.get("target"),
            "Match": "✓" if ca.get("pass") else "⚠",
        })
        
        df_stats = _pd.DataFrame(rows)
        stats_verify_csv = df_stats.to_csv(index=False).encode("utf-8")
        st.download_button(
            "📥 Download Stats Verification CSV",
            data=stats_verify_csv,
            file_name=f"{slug}_stats_verified.csv",
            mime="text/csv",
            use_container_width=True
        )

    meta = {
        "topic": st.session_state.topic,
        "domain": st.session_state.domain_analysis.get("detected_domain",""),
        "tier": st.session_state.tier,
        "stats_level": st.session_state.stats_level,
        "word_count": word_count(paper),
        "citations_verified": len(st.session_state.citation_bank),
        "narrative_selected": st.session_state.selected_narrative,
    }
    meta_bytes = json.dumps(meta, indent=2).encode("utf-8")

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf,"w") as zf:
        zf.writestr(f"{slug}.md", paper)
        zf.writestr(f"{slug}.docx", docx_bytes)
        if csv_bytes:
            zf.writestr(f"{slug}_dataset.csv", csv_bytes)
        zf.writestr(f"{slug}_metadata.json", meta_bytes)

    c1,c2,c3,c4 = st.columns(4)
    c1.download_button("📄 DOCX", docx_bytes, f"{slug}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)
    c2.download_button("📝 Markdown", paper.encode("utf-8"), f"{slug}.md", "text/markdown",
                        use_container_width=True)
    if csv_bytes:
        c3.download_button("📊 Dataset CSV", csv_bytes, f"{slug}_dataset.csv", "text/csv",
                            use_container_width=True)
    c4.download_button("🗜️ ZIP (All)", zip_buf.getvalue(), f"{slug}_complete.zip", "application/zip",
                        use_container_width=True)

    st.markdown("---")
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Words", f"{word_count(paper):,}")
    c2.metric("Citations", len(st.session_state.citation_bank))
    c3.metric("Stats Level", st.session_state.stats_level)
    c4.metric("Narrative", st.session_state.selected_narrative or "-")

    # Supervisor unlock
    st.markdown("---")
    st.markdown("### 👥 Supervisor Review Link — ₹199")
    if st.button("🔓 Unlock Supervisor Link (₹199)"):
        cost = CREDIT_COSTS["supervisor_unlock"]
        ok, msg = credits.reserve(st.session_state.user_id, cost, st.session_state.paper_id)
        if ok:
            show_credit_drain(cost, "supervisor unlock")
            review_token = str(uuid.uuid4())[:8]
            st.markdown(f'<div class="ok-box">✅ Share this link with your supervisor:<br>'
                        f'<strong>https://paperforge.app/review/{review_token}</strong><br>'
                        f'Read-only access. Comments enabled. Valid 30 days.</div>',
                        unsafe_allow_html=True)
        else:
            st.error(msg)

    # Feedback
    st.markdown("---")
    if not st.session_state.feedback_given:
        fb = st.radio("Did this paper meet your expectations?",
                       ["Yes","Partially","No"], horizontal=True)
        if st.button("Submit Feedback"):
            if fb == "Yes":
                bank.positive_feedback(st.session_state.template_key)
            elif fb == "No":
                bank.negative_feedback(st.session_state.template_key)
            st.session_state.feedback_given = True
            st.success("Thank you. Prompt bank updated.")

    if st.button("🆕 Start New Paper", use_container_width=True):
        for k, v in DEFAULTS.items():
            st.session_state[k] = v
        st.session_state.paper_id = str(uuid.uuid4())
        st.rerun()
